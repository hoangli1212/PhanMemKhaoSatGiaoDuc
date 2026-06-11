# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "ThietKeCoSoDuLieu_ToiGian_HeThongKhaoSat.docx"

TABLES = [
    ("users", "Lưu tài khoản, vai trò và nhóm đối tượng của người dùng."),
    ("surveys", "Lưu thông tin khảo sát, người tạo, nhóm đối tượng, thời gian và trạng thái."),
    ("questions", "Lưu câu hỏi thuộc từng khảo sát."),
    ("question_options", "Lưu phương án trả lời cho câu hỏi trắc nghiệm hoặc nhiều lựa chọn."),
    ("responses", "Lưu một lượt gửi phản hồi của người tham gia."),
    ("answers", "Lưu chi tiết câu trả lời cho từng câu hỏi."),
]

RELATIONSHIPS = [
    "Một người dùng có thể tạo nhiều khảo sát.",
    "Một khảo sát có nhiều câu hỏi.",
    "Một câu hỏi có nhiều phương án trả lời.",
    "Một khảo sát có nhiều lượt phản hồi trong responses.",
    "Một lượt phản hồi có nhiều câu trả lời chi tiết trong answers.",
    "Một câu trả lời liên kết đến một câu hỏi và có thể liên kết đến một phương án trả lời.",
]

FLOW = [
    "Admin hoặc cán bộ khảo sát được tạo trong bảng users.",
    "Người tạo khảo sát tạo khảo sát mới trong bảng surveys.",
    "Người tạo khảo sát thêm câu hỏi vào bảng questions.",
    "Nếu câu hỏi có lựa chọn, hệ thống lưu phương án trong question_options.",
    "Người tham gia trả lời khảo sát, hệ thống lưu một bản ghi phản hồi.",
    "Chi tiết từng câu trả lời được lưu vào bảng answers.",
    "Hệ thống thống kê dữ liệu trực tiếp từ responses và answers.",
]


def set_run(run, size=12, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 14 if level == 1 else 12.5, True, (30, 64, 175))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, 12)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(5)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r, 12)
    p.paragraph_format.line_spacing = 1.15


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THIẾT KẾ CƠ SỞ DỮ LIỆU TỐI GIẢN")
    set_run(r, 16, True, (15, 23, 42))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hệ thống khảo sát lấy ý kiến các bên liên quan trong lĩnh vực giáo dục")
    set_run(r, 12, True, (30, 64, 175))

    heading(doc, "1. Mục tiêu thiết kế")
    para(
        doc,
        "Cơ sở dữ liệu được rút gọn để dễ triển khai nhưng vẫn đáp ứng luồng chính: người dùng đăng nhập, "
        "người tạo khảo sát tạo khảo sát, thêm câu hỏi, người tham gia trả lời và hệ thống thống kê kết quả.",
    )

    heading(doc, "2. Danh sách bảng")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["STT", "Tên bảng", "Chức năng"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for idx, (name, desc) in enumerate(TABLES, 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = name
        row[2].text = desc

    heading(doc, "3. Quan hệ dữ liệu")
    for item in RELATIONSHIPS:
        bullet(doc, item)

    heading(doc, "4. Các bảng đã lược bỏ")
    removed = [
        "roles và user_roles: vai trò được lưu trực tiếp bằng trường role trong users.",
        "stakeholder_groups và survey_target_groups: nhóm đối tượng được lưu trực tiếp bằng enum.",
        "question_bank và survey_questions: câu hỏi thuộc trực tiếp một khảo sát.",
        "survey_invitations: chưa quản lý token/lời mời chi tiết ở phiên bản đơn giản.",
        "reports: báo cáo được sinh trực tiếp từ dữ liệu phản hồi, chưa cần lưu lịch sử file.",
    ]
    for item in removed:
        bullet(doc, item)

    heading(doc, "5. Các bảng quan trọng")
    heading(doc, "5.1. users", 2)
    para(doc, "Lưu thông tin tài khoản đăng nhập. Các trường chính gồm id, full_name, email, password_hash, role, stakeholder_group, status.")
    heading(doc, "5.2. surveys", 2)
    para(doc, "Lưu thông tin khảo sát. Các trường chính gồm id, title, description, creator_id, target_group, start_date, end_date, status.")
    heading(doc, "5.3. questions và question_options", 2)
    para(doc, "Lưu câu hỏi thuộc khảo sát và các phương án trả lời. Hệ thống hỗ trợ câu hỏi trắc nghiệm, nhiều lựa chọn, thang điểm và tự luận.")
    heading(doc, "5.4. responses và answers", 2)
    para(doc, "Lưu lượt phản hồi và chi tiết từng câu trả lời. Dashboard và báo cáo có thể thống kê trực tiếp từ hai bảng này.")

    heading(doc, "6. Quy trình dữ liệu chính")
    for item in FLOW:
        bullet(doc, item)

    heading(doc, "7. File triển khai")
    para(doc, "File SQL tạo cơ sở dữ liệu nằm tại: database/schema.sql. Thiết kế hiện phù hợp với MySQL 8+ và sử dụng utf8mb4 để hỗ trợ tiếng Việt.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
